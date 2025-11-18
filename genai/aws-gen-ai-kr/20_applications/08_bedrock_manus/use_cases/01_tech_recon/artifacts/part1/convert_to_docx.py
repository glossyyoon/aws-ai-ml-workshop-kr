from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Load the markdown report
with open("finalreportwithcitations.md", 'r', encoding='utf-8') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Function to add formatted text
def add_paragraph_with_style(text, style='Normal'):
    """Add paragraph with appropriate style"""
    # Remove markdown formatting
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
    text = re.sub(r'`(.+?)`', r'\1', text)  # Code
    
    p = doc.add_paragraph(text, style=style)
    return p

# Parse and convert markdown to docx
lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    # Handle headers
    if line.startswith('# '):
        # H1 - Title
        title = line[2:].strip()
        p = doc.add_heading(title, level=1)
    elif line.startswith('## '):
        # H2 - Major section
        heading = line[3:].strip()
        doc.add_heading(heading, level=2)
    elif line.startswith('### '):
        # H3 - Subsection
        heading = line[4:].strip()
        doc.add_heading(heading, level=3)
    elif line.startswith('#### '):
        # H4 - Technology name
        heading = line[5:].strip()
        doc.add_heading(heading, level=4)
    elif line.startswith('**') and line.endswith('**'):
        # Bold paragraph
        text = line.strip('*')
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
    elif line.startswith('- ') or line.startswith('* '):
        # Bullet point
        text = line[2:].strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        doc.add_paragraph(text, style='List Bullet')
    elif line.startswith('|') and '|' in line:
        # Table row - skip for now (tables are complex in python-docx)
        pass
    elif line.startswith('---'):
        # Horizontal rule - add page break
        doc.add_page_break()
    elif line.startswith('[') and ']' in line:
        # Citation
        text = line.strip()
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
    else:
        # Regular paragraph
        if line:
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text)
    
    i += 1

# Save the document
doc.save("finalreportwithcitations.docx")

print("✅ DOCX conversion complete")
print("   File: finalreportwithcitations.docx")

# Also create a clean version without citations
with open("finalreportwithcitations.md", 'r', encoding='utf-8') as f:
    clean_content = f.read()

# Remove citation markers [1], [2], etc.
clean_content = re.sub(r'\[\d+\]', '', clean_content)

# Save clean markdown
with open("finalreport.md", 'w', encoding='utf-8') as f:
    f.write(clean_content)

# Convert clean version to docx
doc_clean = Document()

# Set margins
sections = doc_clean.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Parse clean content
lines = clean_content.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    if line.startswith('# '):
        doc_clean.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc_clean.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc_clean.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        doc_clean.add_heading(line[5:].strip(), level=4)
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        doc_clean.add_paragraph(text, style='List Bullet')
    elif line.startswith('---'):
        doc_clean.add_page_break()
    elif line and not line.startswith('|'):
        text = line.strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        if text:
            doc_clean.add_paragraph(text)
    
    i += 1

doc_clean.save("finalreport.docx")

print("✅ Clean DOCX conversion complete")
print("   File: finalreport.docx")

